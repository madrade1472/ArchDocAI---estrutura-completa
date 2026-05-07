"""
Layer 3 - Output: Generate a professional .docx technical documentation file.
"""

from pathlib import Path
from dataclasses import dataclass
from ..analysis.analyzer import AnalysisResult
from src.logger import get_logger

log = get_logger(__name__)


def _layer_rationale(layer: dict, all_layers: list, language: str) -> str:
    """Generate a pipeline-position note with information NOT already in the description.

    Shows: what feeds into this layer, what it feeds out to, and deduplicated base
    technology names. Never repeats the description paragraph.
    """
    layer_id = layer.get("id")
    comps = layer.get("components", [])
    conn_ids = layer.get("connections_to", [])
    id_to_name = {lyr["id"]: lyr["name"] for lyr in all_layers}

    # Layers that feed INTO this one (reverse lookup)
    fed_by = [lyr["name"] for lyr in all_layers
              if layer_id and layer_id in lyr.get("connections_to", [])]
    # Layers this one feeds OUT TO
    feeds_to = [id_to_name[c] for c in conn_ids if c in id_to_name]

    # Deduplicate base tech names — splits compound strings like "Python / Selenium / BS4"
    raw_techs = [c["tech"] for c in comps if c.get("tech")]
    base_techs: list[str] = []
    seen: set[str] = set()
    for compound in raw_techs:
        for part in compound.replace(" / ", "/").replace(", ", "/").split("/"):
            normalized = part.strip()
            key = normalized.lower()
            if key and key not in seen:
                seen.add(key)
                base_techs.append(normalized)
    base_techs = base_techs[:4]

    if language == "pt":
        if fed_by and feeds_to:
            flow = f"No pipeline, recebe dados de {', '.join(fed_by)} e entrega para {', '.join(feeds_to)}."
        elif fed_by:
            flow = f"Recebe dados de {', '.join(fed_by)} e representa o estágio final do pipeline."
        elif feeds_to:
            flow = f"É o ponto de entrada do pipeline e alimenta diretamente {', '.join(feeds_to)}."
        else:
            flow = "Opera de forma independente, sem dependências diretas de outras camadas."
        tech_str = f" Tecnologias-chave: {', '.join(base_techs)}." if base_techs else ""
        return f"↳ {flow}{tech_str}"
    else:
        if fed_by and feeds_to:
            flow = f"In the pipeline, it receives data from {', '.join(fed_by)} and delivers to {', '.join(feeds_to)}."
        elif fed_by:
            flow = f"Receives data from {', '.join(fed_by)} — final stage of the pipeline."
        elif feeds_to:
            flow = f"Pipeline entry point, feeding directly into {', '.join(feeds_to)}."
        else:
            flow = "Operates independently, with no direct dependencies on other layers."
        tech_str = f" Key technologies: {', '.join(base_techs)}." if base_techs else ""
        return f"↳ {flow}{tech_str}"


@dataclass
class DocxGenerator:
    output_dir: str = "./output"
    language: str = "pt"

    def generate(self, result: AnalysisResult, diagram_path: str | None = None,
                 interactive_diagram_path: str | None = None) -> str:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import docx.oxml

        doc = Document()

        # ── Styles ──────────────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Cover ────────────────────────────────────────────────────────────
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(result.project_name)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

        subtitle_label = "Documentação Técnica de Arquitetura" if self.language == "pt" else "Technical Architecture Documentation"
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(subtitle_label)
        sub_run.italic = True
        sub_run.font.size = Pt(14)

        doc.add_paragraph()

        # ── Section helper ───────────────────────────────────────────────────
        def add_section(title: str):
            h = doc.add_heading(title, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

        def add_subsection(title: str):
            h = doc.add_heading(title, level=2)
            h.runs[0].font.color.rgb = RGBColor(0x45, 0x78, 0x9B)

        def add_bullet(text: str):
            p = doc.add_paragraph(text, style="List Bullet")
            p.runs[0].font.size = Pt(11)

        def add_info_block(text: str):
            """Render a light-blue callout box with a blue left border."""
            from docx.oxml import OxmlElement
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()

            # Left border (thick blue bar)
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "thick")
            left.set(qn("w:sz"), "24")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "2563EB")
            pBdr.append(left)
            pPr.append(pBdr)

            # Light-blue background shading
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "DBEAFE")
            pPr.append(shd)

            # Indentation
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "200")
            pPr.append(ind)

            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            doc.add_paragraph()  # breathing space after block

        # ── 1. Overview ──────────────────────────────────────────────────────
        section_title = "1. Visão Geral do Projeto" if self.language == "pt" else "1. Project Overview"
        add_section(section_title)
        doc.add_paragraph(result.description)

        tech_title = "Stack Tecnológico" if self.language == "pt" else "Technology Stack"
        add_subsection(tech_title)
        for tech in result.tech_stack:
            add_bullet(tech)

        # ── 2. Architecture Layers ───────────────────────────────────────────
        arch_title = "2. Arquitetura em Camadas" if self.language == "pt" else "2. Layered Architecture"
        add_section(arch_title)

        for layer in result.layers:
            add_subsection(f"{layer['name']}")
            doc.add_paragraph(layer.get("description", ""))

            # Info block — design rationale for this layer
            add_info_block(_layer_rationale(layer, result.layers, self.language))

            components = layer.get("components", [])
            if components:
                comp_label = "Componentes:" if self.language == "pt" else "Components:"
                doc.add_paragraph(comp_label)
                for comp in components:
                    tech_note = f" ({comp['tech']})" if comp.get("tech") else ""
                    add_bullet(f"{comp['name']}{tech_note} - {comp.get('description', '')}")

        # ── 3. Diagram ───────────────────────────────────────────────────────
        has_any_diagram = (diagram_path and Path(diagram_path).exists()) or \
                          (interactive_diagram_path and Path(interactive_diagram_path).exists())
        if has_any_diagram:
            diag_title = "3. Diagrama da Arquitetura" if self.language == "pt" else "3. Architecture Diagram"
            add_section(diag_title)

            if diagram_path and Path(diagram_path).exists():
                doc.add_picture(diagram_path, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if interactive_diagram_path:
                p = Path(interactive_diagram_path)
                if not p.exists():
                    log.warning("DOCX: interactive_diagram_path provided but file missing: %s", p)
                elif p.stat().st_size < 1024:
                    log.warning("DOCX: interactive_diagram_path too small (%d bytes), skipping: %s", p.stat().st_size, p)
                else:
                    sub_title = "Diagrama Interativo (Node-Graph)" if self.language == "pt" else "Interactive Diagram (Node-Graph)"
                    add_subsection(sub_title)
                    try:
                        doc.add_picture(str(p), width=Inches(6.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as exc:
                        log.error("DOCX: failed to embed interactive PNG at %s: %s", p, exc, exc_info=True)
            else:
                log.info("DOCX: no interactive_diagram_path provided - skipping interactive subsection")

        # ── 4. Quality Score ─────────────────────────────────────────────────
        next_n = 4 if has_any_diagram else 3
        score = result.quality_score or {}
        if score.get("total"):
            qs_title = f"{next_n}. Score de Qualidade Arquitetural" if self.language == "pt" else f"{next_n}. Architecture Quality Score"
            add_section(qs_title)

            total = score.get("total", 0)
            label_pt = ("Excelente" if total >= 80 else
                        "Bom, com pontos a evoluir" if total >= 60 else
                        "Funcional mas com lacunas relevantes" if total >= 40 else
                        "Atencao: problemas estruturais")
            label_en = ("Excellent" if total >= 80 else
                        "Good, with room to improve" if total >= 60 else
                        "Functional but with relevant gaps" if total >= 40 else
                        "Warning: structural problems")
            label = label_pt if self.language == "pt" else label_en

            score_p = doc.add_paragraph()
            r1 = score_p.add_run(f"{total}/100 ")
            r1.bold = True
            r1.font.size = Pt(20)
            r1.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
            r2 = score_p.add_run(f"  {label}")
            r2.italic = True
            r2.font.size = Pt(11)
            r2.font.color.rgb = RGBColor(0x45, 0x78, 0x9B)

            if score.get("rationale"):
                doc.add_paragraph(score["rationale"])

            breakdown = score.get("breakdown") or {}
            dim_labels_pt = {"arquitetura": "Arquitetura", "codigo": "Codigo",
                             "documentacao": "Documentacao", "testabilidade": "Testabilidade",
                             "devops": "DevOps"}
            dim_labels_en = {"arquitetura": "Architecture", "codigo": "Code",
                             "documentacao": "Documentation", "testabilidade": "Testability",
                             "devops": "DevOps"}
            dim_labels = dim_labels_pt if self.language == "pt" else dim_labels_en
            for k, dim_name in dim_labels.items():
                v = int(breakdown.get(k, 0))
                add_bullet(f"{dim_name}: {v}/20")

            next_n += 1

        # ── Architecture Pattern (classification badge) ──────────────────────
        pattern = getattr(result, "architecture_pattern", None) or {}
        matches = pattern.get("matches") or []
        if matches:
            ap_title = (f"{next_n}. Padrao Arquitetural"
                        if self.language == "pt"
                        else f"{next_n}. Architectural Pattern")
            add_section(ap_title)

            primary = pattern.get("primary") or matches[0]["name"]
            primary_pct = matches[0].get("adherence", 0)
            badge_p = doc.add_paragraph()
            r1 = badge_p.add_run(f"{primary}  ")
            r1.bold = True
            r1.font.size = Pt(18)
            r1.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
            r2 = badge_p.add_run(f"{primary_pct}%")
            r2.bold = True
            r2.font.size = Pt(18)
            r2.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

            if pattern.get("summary"):
                doc.add_paragraph(pattern["summary"])

            other_lbl = ("Outras correspondencias:"
                         if self.language == "pt" else
                         "Other matches:")
            if len(matches) > 1:
                doc.add_paragraph(other_lbl)

            for m in matches:
                title_p = doc.add_paragraph()
                pr1 = title_p.add_run(f"{m['name']} - {m.get('adherence', 0)}%")
                pr1.bold = True
                pr1.font.size = Pt(11)
                if m.get("rationale"):
                    doc.add_paragraph(m["rationale"])
                ev_lbl = "Evidencias:" if self.language == "pt" else "Evidence:"
                evidence = m.get("evidence") or []
                if evidence:
                    el = doc.add_paragraph()
                    el.add_run(ev_lbl).bold = True
                    for e in evidence:
                        add_bullet(e)
            next_n += 1

        # ── Use Cases (sequence diagrams) ────────────────────────────────────
        use_cases = getattr(result, "use_cases", None) or []
        if use_cases:
            from .mermaid_renderer import render_mermaid_png
            uc_title = f"{next_n}. Diagramas de Sequencia" if self.language == "pt" else f"{next_n}. Sequence Diagrams"
            add_section(uc_title)
            for uc in use_cases:
                add_subsection(uc.get("name", ""))
                if uc.get("description"):
                    doc.add_paragraph(uc["description"])
                diagram = (uc.get("sequence_diagram") or "").strip()
                if not diagram:
                    continue
                # Render via kroki.io (cached). Fall back to monospace text
                # if the service is unreachable or rejects the diagram.
                png_path = render_mermaid_png(diagram)
                if png_path:
                    doc.add_picture(str(png_path), width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    log.warning("DOCX: rendering failed, embedding text for use case '%s'", uc.get("name", ""))
                    code_p = doc.add_paragraph()
                    code_run = code_p.add_run(diagram)
                    code_run.font.name = "Consolas"
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
            next_n += 1

        # ── Architecture Decision Records (ADRs) ─────────────────────────────
        adrs = getattr(result, "adrs", None) or []
        if adrs:
            adr_title = (f"{next_n}. Decisoes Arquiteturais (ADRs)"
                         if self.language == "pt"
                         else f"{next_n}. Architecture Decision Records (ADRs)")
            add_section(adr_title)
            intro = (
                "Decisoes arquiteturais identificadas no projeto. Cada ADR descreve o contexto, "
                "a decisao tomada e suas consequencias — formato MADR."
                if self.language == "pt" else
                "Architectural decisions identified in the project. Each ADR describes the "
                "context, the decision and its consequences — MADR format."
            )
            doc.add_paragraph(intro)

            for i, adr in enumerate(adrs, start=1):
                num = f"{i:04d}"
                title = (adr.get("title") or "Untitled").strip()
                status = (adr.get("status") or "accepted").capitalize()
                add_subsection(f"ADR-{num}: {title}")

                status_p = doc.add_paragraph()
                status_label = "Status:" if self.language == "pt" else "Status:"
                run_lbl = status_p.add_run(status_label + " ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(11)
                status_p.add_run(status).font.size = Pt(11)

                def _adr_field(label: str, value: str):
                    if not value:
                        return
                    p = doc.add_paragraph()
                    r = p.add_run(label + " ")
                    r.bold = True
                    r.font.size = Pt(11)
                    p.add_run(value).font.size = Pt(11)

                ctx_lbl = "Contexto:" if self.language == "pt" else "Context:"
                dec_lbl = "Decisao:" if self.language == "pt" else "Decision:"
                cons_lbl = "Consequencias:" if self.language == "pt" else "Consequences:"
                alt_lbl = "Alternativas:" if self.language == "pt" else "Alternatives:"
                _adr_field(ctx_lbl, (adr.get("context") or "").strip())
                _adr_field(dec_lbl, (adr.get("decision") or "").strip())
                _adr_field(cons_lbl, (adr.get("consequences") or "").strip())
                _adr_field(alt_lbl, (adr.get("alternatives") or "").strip())

            next_n += 1

        # ── Good Practices ───────────────────────────────────────────────────
        gp_title = f"{next_n}. Boas Práticas Identificadas" if self.language == "pt" else f"{next_n}. Good Practices Identified"
        add_section(gp_title)
        for gp in result.good_practices:
            add_bullet(gp)

        # ── Improvement Points ───────────────────────────────────────────────
        next_n += 1
        ip_title = f"{next_n}. Pontos de Melhoria" if self.language == "pt" else f"{next_n}. Improvement Points"
        add_section(ip_title)
        for ip in result.improvement_points:
            add_bullet(ip)

        # ── User Corrections (if any) ────────────────────────────────────────
        if result.user_corrections:
            next_n += 1
            uc_title = f"{next_n}. Ajustes Validados pelo Usuário" if self.language == "pt" else f"{next_n}. User-Validated Adjustments"
            add_section(uc_title)
            for uc in result.user_corrections:
                add_bullet(uc)

        # ── Save ─────────────────────────────────────────────────────────────
        out_path = Path(self.output_dir) / f"{self._safe_name(result.project_name)}_architecture.docx"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return str(out_path)

    def _safe_name(self, name: str) -> str:
        return "".join(c if c.isalnum() or c in "-_" else "_" for c in name)
